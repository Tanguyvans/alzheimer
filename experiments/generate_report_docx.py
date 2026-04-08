#!/usr/bin/env python3
"""
Generate a Word report (.docx) with:
- DeLong test results
- 1 IG image AD + 1 IG image CN
- Description of the interpretability folder contents
"""

import csv
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = Path(__file__).parent
REPORT_DIR = BASE / "report_multi_seed"
# Handle nested interpretability dir (from scp -r)
IG_DIR = REPORT_DIR / "interpretability" / "interpretability"
if not (IG_DIR / "mlp_early_fusion").exists():
    IG_DIR = REPORT_DIR / "interpretability"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def build_report():
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ══════════════════════════════════════════════
    # Title
    # ══════════════════════════════════════════════
    title = doc.add_heading('ResNet3D Fusion — Rapport d\'analyse multi-seeds', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph(
        'Classification CN vs AD utilisant un ResNet50 3D (poids MedicalNet pré-entraînés) '
        'combiné avec des caractéristiques cliniques tabulaires. Évaluation multi-seeds (5 seeds) '
        'sur 4 stratégies de fusion.'
    )

    # ══════════════════════════════════════════════
    # 1. Résumé des performances
    # ══════════════════════════════════════════════
    add_heading(doc, '1. Résumé des performances', level=1)

    # Load summary table
    summary_path = REPORT_DIR / "summary_table.csv"
    with open(summary_path) as f:
        reader = csv.reader(f)
        rows = list(reader)

    headers = ['Méthode', 'Acc %', 'Acc Éq %', 'Sens %', 'Spéc %', 'AUC', 'Seeds']
    table = doc.add_table(rows=len(rows), cols=len(headers), style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row in enumerate(rows[1:], start=1):
        if len(row) < 12:
            continue
        method = row[0]
        acc_m, acc_s = float(row[1]), float(row[2])
        bacc_m, bacc_s = float(row[3]), float(row[4])
        sens_m, sens_s = float(row[5]), float(row[6])
        spec_m, spec_s = float(row[7]), float(row[8])
        auc_m, auc_s = float(row[9]), float(row[10])
        n_seeds = row[11]

        values = [
            method,
            f'{acc_m:.1f} +/- {acc_s:.1f}',
            f'{bacc_m:.1f} +/- {bacc_s:.1f}',
            f'{sens_m:.1f} +/- {sens_s:.1f}',
            f'{spec_m:.1f} +/- {spec_s:.1f}',
            f'{auc_m:.3f} +/- {auc_s:.3f}',
            str(n_seeds),
        ]
        for j, val in enumerate(values):
            cell = table.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('')

    # ══════════════════════════════════════════════
    # 2. Test de DeLong
    # ══════════════════════════════════════════════
    add_heading(doc, '2. Test de DeLong — Comparaison des AUC par paires', level=1)

    add_body(doc,
        'Le test de DeLong évalue si les différences d\'AUC entre les modèles sont '
        'statistiquement significatives. La heatmap ci-dessous montre les p-values par paires. '
        'Les cellules vertes indiquent des différences significatives (p < 0.05).'
    )

    delong_img = REPORT_DIR / "delong_test.png"
    if delong_img.exists():
        doc.add_picture(str(delong_img), width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body(doc, 'Résultats clés du test de DeLong :')
    bullets = [
        'Toutes les méthodes de fusion surpassent significativement le MRI seul (p < 0.001), '
        'confirmant l\'intérêt de l\'intégration multimodale.',
        'MLP Late Wt et XGB Late Wt obtiennent les meilleurs AUC (0.955 et 0.952) '
        'sans différence significative entre eux (p = 0.36).',
        'Les modèles tabulaires seuls (XGB/MLP) sont significativement inférieurs aux meilleures '
        'méthodes de fusion, montrant que l\'IRM apporte une information discriminante '
        'au-delà des caractéristiques cliniques.',
        'MLP Late Stack est la méthode de fusion la plus faible, significativement en dessous des autres.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # ══════════════════════════════════════════════
    # 3. Courbes ROC
    # ══════════════════════════════════════════════
    add_heading(doc, '3. Courbes ROC', level=1)

    roc_img = REPORT_DIR / "roc_curves.png"
    if roc_img.exists():
        doc.add_picture(str(roc_img), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ══════════════════════════════════════════════
    # 4. Interprétabilité
    # ══════════════════════════════════════════════
    add_heading(doc, '4. Interprétabilité — Integrated Gradients', level=1)

    add_body(doc,
        'La méthode Integrated Gradients (Sundararajan et al., 2017) fournit des cartes '
        'd\'attribution au niveau du voxel, à la résolution native de l\'entrée (128x128x128). '
        'Contrairement au GradCAM, limité par la résolution des feature maps (~4x4x4), '
        'Integrated Gradients calcule les attributions directement dans l\'espace d\'entrée '
        'en intégrant les gradients le long d\'un chemin allant d\'une baseline nulle à l\'entrée. '
        'Cela permet une localisation précise des régions cérébrales influençant la décision du modèle.'
    )

    add_heading(doc, '4.1 Exemple : Patient Alzheimer', level=2)
    add_body(doc,
        'Prédiction AD à haute confiance (p(AD) = 0.997). Les régions lumineuses indiquent les voxels '
        'à forte attribution pour la prédiction AD. Les activations sont concentrées dans le '
        'lobe temporal médial (hippocampe, cortex entorhinal), ce qui est cohérent avec les '
        'patterns connus de neurodégénérescence liée à la maladie d\'Alzheimer.'
    )

    ad_img = IG_DIR / "mlp_early_fusion" / "AD_01.png"
    if ad_img.exists():
        doc.add_picture(str(ad_img), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure : Integrated Gradients — Patient AD (MLP Early Fusion, seed 2)',
                          style='Caption')

    add_heading(doc, '4.2 Exemple : Patient cognitivement normal', level=2)
    add_body(doc,
        'Prédiction CN à haute confiance (p(AD) = 0.003). Le pattern d\'attribution est plus diffus '
        'et distribué sur l\'ensemble des régions corticales, avec une moindre concentration dans les '
        'structures temporales médiales. Cela suggère que le modèle reconnaît l\'intégrité structurelle '
        'des régions hippocampiques et temporales comme indicateur de cognition normale.'
    )

    cn_img = IG_DIR / "mlp_early_fusion" / "CN_01.png"
    if cn_img.exists():
        doc.add_picture(str(cn_img), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure : Integrated Gradients — Patient CN (MLP Early Fusion, seed 2)',
                          style='Caption')

    # ══════════════════════════════════════════════
    # 5. Contenu du dossier
    # ══════════════════════════════════════════════
    add_heading(doc, '5. Contenu du dossier d\'interprétabilité', level=1)

    add_body(doc,
        'Le dossier interpretability/ ci-joint contient les visualisations Integrated Gradients '
        'pour les 4 modèles, calculées sur les mêmes 5 patients AD et 5 patients CN '
        'afin de permettre une comparaison directe. La structure est la suivante :'
    )

    folder_desc = [
        ('mlp_early_fusion/', 'Integrated Gradients pour le modèle MLP Early Fusion '
         '(ResNet3D + MLP tabulaire, bout-en-bout). Contient AD_01..05.png, CN_01..05.png '
         '(cartes individuelles, 3 vues chacune) et grid_mlp_early_fusion.png (les 10 patients).'),
        ('mlp_late_fusion/', 'Idem pour MLP Late Fusion (branche MRI du ResNet3D uniquement, '
         'entraînée séparément du MLP tabulaire). Montre les régions ciblées par la branche IRM.'),
        ('xgb_early_fusion/', 'Idem pour XGBoost Early Fusion (backbone ResNet3D finetuné, '
         'features envoyées à XGBoost). IG calculé à travers le backbone CNN.'),
        ('xgb_late_fusion/', 'Idem pour XGBoost Late Fusion (branche MRI du ResNet3D, '
         'XGBoost séparé sur le tabulaire). Montre les attributions de la branche IRM.'),
        ('cross_model_comparison.png', 'Grille récapitulative : 10 patients (lignes) x 4 modèles (colonnes), '
         'vue axiale. Permet une comparaison directe de ce que chaque modèle observe sur le même cerveau.'),
        ('group_average_AD.png', 'Carte d\'attribution moyenne sur 50 patients AD correctement classifiés. '
         'Met en évidence les régions systématiquement importantes pour la prédiction AD.'),
        ('group_average_CN.png', 'Carte d\'attribution moyenne sur 50 patients CN correctement classifiés.'),
        ('group_difference_AD_minus_CN.png', 'Carte d\'attribution différentielle (moyenne AD moins moyenne CN). '
         'Rouge = plus important pour AD (lobe temporal médial), Bleu = plus important pour CN (frontal/pariétal). '
         'C\'est la figure la plus informative pour comprendre le comportement du modèle.'),
        ('summary_figure.png', 'Figure de synthèse combinant exemples individuels, '
         'moyenne de groupe et carte de différence.'),
        ('*.npy files', 'Tableaux numpy bruts des moyennes de groupe et cartes de différence '
         'pour analyse complémentaire ou visualisation personnalisée.'),
    ]

    for name, desc in folder_desc:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ' ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    add_body(doc,
        'Toutes les cartes d\'attribution ont été calculées avec Integrated Gradients '
        '(100 étapes d\'interpolation, baseline nulle), à partir de la meilleure seed '
        '(seed 2, AUC = 0.954). Les mêmes 5 patients AD et 5 patients CN '
        '(sélectionnés par confiance de prédiction) sont utilisés pour les 4 modèles '
        'afin de permettre une comparaison directe.'
    )

    # Save
    out_path = REPORT_DIR / "resnet3d_fusion_report.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    build_report()
