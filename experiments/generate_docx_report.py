#!/usr/bin/env python3
"""
Generate a .docx report from multi-seed results, matching the style
of report_old/resnet3d_fusion_report.docx. Report in French.
"""

import numpy as np
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from sklearn.metrics import roc_auc_score, accuracy_score, balanced_accuracy_score, confusion_matrix

BASE = Path(__file__).parent
MLP_DIR = BASE / "resnet3d_mlp"
XGB_DIR = BASE / "resnet3d_xgboost"
REPORT_DIR = BASE / "report_multi_seed"
REPORT_DIR.mkdir(exist_ok=True)

MAX_SEED = 20
TARGET_SEEDS = 5  # expected number of seeds per method


def load_all_seeds():
    all_preds = {}
    seed_counts = {}
    y_true = None

    def _add(method, seed, proba, yt):
        nonlocal y_true
        y_true = yt
        if method not in all_preds:
            all_preds[method] = []
            seed_counts[method] = []
        if len(all_preds[method]) >= TARGET_SEEDS:
            return  # cap at TARGET_SEEDS
        all_preds[method].append(proba)
        seed_counts[method].append(seed)

    for seed in range(MAX_SEED):
        p = MLP_DIR / "results_early" / f"seed_{seed}"
        if (p / "y_true_test.npy").exists():
            yt = np.load(p / "y_true_test.npy")
            _add("MLP Early", seed, np.load(p / "y_proba_test.npy"), yt)

    for seed in range(MAX_SEED):
        p = MLP_DIR / "results_late_fusion" / f"seed_{seed}"
        if (p / "y_true_test.npy").exists():
            yt = np.load(p / "y_true_test.npy")
            # MRI only (MLP) skipped — redundant with MRI only
            _add("Tab only (MLP)", seed, np.load(p / "y_proba_tab_test.npy"), yt)
            _add("MLP Late Avg", seed, np.load(p / "y_proba_avg_test.npy"), yt)
            _add("MLP Late Wt", seed, np.load(p / "y_proba_weighted_test.npy"), yt)
            _add("MLP Late Stack", seed, np.load(p / "y_proba_stacking_test.npy"), yt)

    for seed in range(MAX_SEED):
        p = XGB_DIR / "results_finetuned" / f"seed_{seed}"
        if (p / "y_true_test.npy").exists():
            yt = np.load(p / "y_true_test.npy")
            _add("XGB Early", seed, np.load(p / "y_proba_test.npy"), yt)

    for seed in range(MAX_SEED):
        p = XGB_DIR / "results_late_fusion" / f"seed_{seed}"
        if (p / "y_true_test.npy").exists():
            yt = np.load(p / "y_true_test.npy")
            _add("MRI only", seed, np.load(p / "y_proba_mri_test.npy"), yt)
            _add("Tab only (XGB)", seed, np.load(p / "y_proba_tab_test.npy"), yt)
            _add("XGB Late Avg", seed, np.load(p / "y_proba_avg_test.npy"), yt)
            _add("XGB Late Wt", seed, np.load(p / "y_proba_weighted_test.npy"), yt)
            _add("XGB Late Stack", seed, np.load(p / "y_proba_stacking_test.npy"), yt)

    return y_true, all_preds, seed_counts


def compute_metrics(y_true, all_preds, seed_counts):
    """Compute mean +/- std for each method."""
    rows = []

    method_info = [
        ("MRI only",        "IRM seule (ResNet3D)",           "\u2014"),
        ("Tab only (XGB)",  "Tabulaire seul (XGBoost)",       "\u2014"),
        ("Tab only (MLP)",  "Tabulaire seul (MLP)",           "\u2014"),
        ("MLP Early",       "ResNet3D + MLP concat",          "Early"),
        ("XGB Early",       "ResNet3D emb + Tab \u2192 XGB",  "Early"),
        ("XGB Late Avg",    "ResNet3D + XGBoost (Moy.)",     "Late"),
        ("XGB Late Wt",     "ResNet3D + XGBoost (Pond.)",    "Late"),
        ("XGB Late Stack",  "ResNet3D + XGBoost (Stacking)", "Late"),
        ("MLP Late Avg",    "ResNet3D + MLP (Moy.)",         "Late"),
        ("MLP Late Wt",     "ResNet3D + MLP (Pond.)",        "Late"),
        ("MLP Late Stack",  "ResNet3D + MLP (Stacking)",     "Late"),
    ]

    for key, display_name, fusion_type in method_info:
        if key not in all_preds:
            continue
        proba_list = all_preds[key]
        n_seeds = len(proba_list)

        accs, baccs, senss, specs, aucs = [], [], [], [], []
        for y_proba in proba_list:
            y_pred = (y_proba >= 0.5).astype(int)
            cm = confusion_matrix(y_true, y_pred)
            tn, fp, fn, tp = cm.ravel()
            accs.append(accuracy_score(y_true, y_pred) * 100)
            baccs.append(balanced_accuracy_score(y_true, y_pred) * 100)
            senss.append(tp / (tp + fn) * 100 if (tp + fn) > 0 else 0)
            specs.append(tn / (tn + fp) * 100 if (tn + fp) > 0 else 0)
            aucs.append(roc_auc_score(y_true, y_proba))

        # AUC ensemble = AUC(mean probabilities) — same as DeLong heatmap
        mean_proba = np.mean(proba_list, axis=0)
        auc_ensemble = roc_auc_score(y_true, mean_proba)

        rows.append({
            'key': key,
            'Method': display_name,
            'Fusion': fusion_type,
            'N': n_seeds,
            'complete': n_seeds >= TARGET_SEEDS,
            'Acc_mean': np.mean(accs), 'Acc_std': np.std(accs) if n_seeds > 1 else None,
            'BAcc_mean': np.mean(baccs), 'BAcc_std': np.std(baccs) if n_seeds > 1 else None,
            'Sens_mean': np.mean(senss), 'Sens_std': np.std(senss) if n_seeds > 1 else None,
            'Spec_mean': np.mean(specs), 'Spec_std': np.std(specs) if n_seeds > 1 else None,
            'AUC_mean': np.mean(aucs), 'AUC_std': np.std(aucs) if n_seeds > 1 else None,
            'AUC_ensemble': auc_ensemble,
        })

    return rows


def fmt_pct(mean, std, complete):
    """Format percentage metric. Show N/A for incomplete single-seed methods."""
    if not complete and std is None:
        return f"{mean:.1f}*"
    if std is None:
        return f"{mean:.1f}"
    return f"{mean:.1f} \u00b1 {std:.1f}"


def fmt_auc(mean, std, complete):
    """Format AUC metric."""
    if not complete and std is None:
        return f"{mean:.3f}*"
    if std is None:
        return f"{mean:.3f}"
    return f"{mean:.3f} \u00b1 {std:.3f}"


def set_cell_font(cell, text, bold=False, size=8):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_report():
    print("Loading predictions...")
    y_true, all_preds, seed_counts = load_all_seeds()
    metrics = compute_metrics(y_true, all_preds, seed_counts)

    # Find best per column (only among complete methods)
    complete_metrics = [r for r in metrics if r['complete']]
    if not complete_metrics:
        complete_metrics = metrics  # fallback if none complete
    best_acc = max(r['Acc_mean'] for r in complete_metrics)
    best_bacc = max(r['BAcc_mean'] for r in complete_metrics)
    best_sens = max(r['Sens_mean'] for r in complete_metrics)
    best_spec = max(r['Spec_mean'] for r in complete_metrics)
    best_auc = max(r['AUC_ensemble'] for r in complete_metrics)

    doc = Document()

    # ── Titre ──
    doc.add_heading("Fusion Multimodale ResNet3D", level=1)
    doc.add_paragraph(
        "Classification CN vs MA \u2014 R\u00e9sultats sur l'ensemble de test (multi-seed)"
    )

    # ── Jeu de donn\u00e9es ──
    doc.add_heading("Jeu de donn\u00e9es", level=2)
    doc.add_paragraph(
        "Jeu de donn\u00e9es combin\u00e9 de trajectoire (ADNI + OASIS + NACC)"
    )
    doc.add_paragraph(
        "Train : 4 245 / Val : 910 / Test : 910 \u00e9chantillons (78% CN / 22% MA)",
        style="List Bullet"
    )
    doc.add_paragraph(
        "16 variables tabulaires : d\u00e9mographiques, tests cognitifs, ant\u00e9c\u00e9dents m\u00e9dicaux",
        style="List Bullet"
    )
    doc.add_paragraph(
        "IRM : volumes c\u00e9r\u00e9braux 3D (128 \u00d7 128 \u00d7 128)",
        style="List Bullet"
    )

    # ── Backbone ──
    doc.add_heading("Backbone : ResNet3D", level=2)
    doc.add_paragraph(
        "MONAI ResNet50 3D pr\u00e9-entra\u00een\u00e9 sur 23 jeux de donn\u00e9es d'imagerie m\u00e9dicale (MedicalNet). "
        "Produit des vecteurs de caract\u00e9ristiques de dimension 2048 \u00e0 partir de volumes IRM 3D.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Fine-tuning : backbone gel\u00e9 pendant les 3 premi\u00e8res \u00e9poques, puis d\u00e9gel\u00e9 avec "
        "taux d'apprentissage diff\u00e9rentiel (backbone 10x plus faible). Pr\u00e9cision mixte (AMP).",
        style="List Bullet"
    )

    # ── Strat\u00e9gies de fusion ──
    doc.add_heading("Strat\u00e9gies de fusion", level=2)

    doc.add_heading("1. Fusion pr\u00e9coce : ResNet3D + MLP", level=3)
    doc.add_paragraph(
        "Les deux modalit\u00e9s sont encod\u00e9es en vecteurs de caract\u00e9ristiques, concat\u00e9n\u00e9s, "
        "puis aliment\u00e9s \u00e0 un classifieur MLP conjoint. Le mod\u00e8le entier est entra\u00een\u00e9 de bout en bout."
    )
    doc.add_paragraph("Branche IRM : ResNet3D (MedicalNet) \u2192 2048 dimensions", style="List Bullet")
    doc.add_paragraph("Branche tabulaire : encodeur MLP [64, 32] avec LayerNorm", style="List Bullet")
    doc.add_paragraph(
        "Fusion : concat\u00e9nation (2080-d) \u2192 MLP [256, 128] \u2192 classification",
        style="List Bullet"
    )

    doc.add_heading("2. Fusion pr\u00e9coce : embeddings ResNet3D + XGBoost", level=3)
    doc.add_paragraph(
        "Le ResNet3D est d'abord fine-tun\u00e9 sur la classification IRM. Ensuite, les embeddings de "
        "dimension 2048 sont extraits et concat\u00e9n\u00e9s avec les 16 variables tabulaires. "
        "Un mod\u00e8le XGBoost unique est entra\u00een\u00e9 sur le vecteur combin\u00e9 de 2064 dimensions."
    )
    doc.add_paragraph(
        "Phase 1 : fine-tuning du ResNet3D avec t\u00eate lin\u00e9aire (30 \u00e9poques)",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Phase 2 : extraction des embeddings 2048-d + 16 variables tabulaires \u2192 XGBoost",
        style="List Bullet"
    )

    doc.add_heading("3. Fusion tardive : pr\u00e9dictions s\u00e9par\u00e9es + combinaison de probabilit\u00e9s", level=3)
    doc.add_paragraph(
        "Chaque modalit\u00e9 produit sa propre pr\u00e9diction ind\u00e9pendante. La branche IRM (ResNet3D fine-tun\u00e9 "
        "avec t\u00eate lin\u00e9aire) produit P(MA|IRM), et la branche tabulaire (MLP ou XGBoost) produit "
        "P(MA|tabulaire). Les deux probabilit\u00e9s sont combin\u00e9es via :"
    )
    doc.add_paragraph("Moyenne : moyenne simple des probabilit\u00e9s", style="List Bullet")
    doc.add_paragraph(
        "Moyenne pond\u00e9r\u00e9e : poids optimis\u00e9s sur l'ensemble de validation",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Stacking : r\u00e9gression logistique entra\u00een\u00e9e sur les probabilit\u00e9s des deux branches",
        style="List Bullet"
    )

    # ── R\u00e9sultats ──
    doc.add_heading("R\u00e9sultats", level=2)

    # Figures
    for img_name, caption in [
        ("boxplots.png", "Figure : Distribution de l'AUC et de la Balanced Accuracy sur les diff\u00e9rents seeds."),
        ("roc_curves.png", "Figure : Courbes ROC (moyenne \u00b1 \u00e9cart-type) pour toutes les m\u00e9thodes."),
    ]:
        img_path = REPORT_DIR / img_name
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.0))
            doc.add_paragraph(caption)

    # ── Tableau des r\u00e9sultats ──
    n_complete = sum(1 for r in metrics if r['complete'])
    n_incomplete = sum(1 for r in metrics if not r['complete'])

    caption_text = (
        f"Tableau : R\u00e9sultats sur l'ensemble de test pour toutes les m\u00e9thodes de fusion "
        f"multimodale ResNet3D (moyenne \u00b1 \u00e9cart-type sur {TARGET_SEEDS} seeds). "
        f"L'AUC est calcul\u00e9e sur les probabilit\u00e9s moyenn\u00e9es (ensemble), "
        f"ce qui correspond aux valeurs du test de DeLong. "
        f"Les valeurs en gras indiquent la meilleure performance par m\u00e9trique."
    )
    if n_incomplete > 0:
        caption_text += (
            f" Les m\u00e9thodes marqu\u00e9es d'un * n'ont qu'un seul seed "
            f"(entra\u00eenement en cours)."
        )
    doc.add_paragraph(caption_text)

    headers = ["M\u00e9thode", "Fusion", "Acc (%)", "Bal Acc (%)", "Sens (%)", "Spec (%)", "AUC"]
    table = doc.add_table(rows=1 + len(metrics), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[j], h, bold=True, size=8)

    # Data rows
    for i, row in enumerate(metrics):
        cells = table.rows[i + 1].cells
        method_label = row['Method']
        if not row['complete']:
            method_label += f" (1 seed)*"
        set_cell_font(cells[0], method_label, size=8)
        set_cell_font(cells[1], row['Fusion'], size=8)

        # Acc
        text = fmt_pct(row['Acc_mean'], row['Acc_std'], row['complete'])
        bold = (row['Acc_mean'] == best_acc and row['complete'])
        set_cell_font(cells[2], text, bold=bold, size=8)

        # Bal Acc
        text = fmt_pct(row['BAcc_mean'], row['BAcc_std'], row['complete'])
        bold = (row['BAcc_mean'] == best_bacc and row['complete'])
        set_cell_font(cells[3], text, bold=bold, size=8)

        # Sens
        text = fmt_pct(row['Sens_mean'], row['Sens_std'], row['complete'])
        bold = (row['Sens_mean'] == best_sens and row['complete'])
        set_cell_font(cells[4], text, bold=bold, size=8)

        # Spec
        text = fmt_pct(row['Spec_mean'], row['Spec_std'], row['complete'])
        bold = (row['Spec_mean'] == best_spec and row['complete'])
        set_cell_font(cells[5], text, bold=bold, size=8)

        # AUC (ensemble = AUC of mean probabilities)
        text = f"{row['AUC_ensemble']:.3f}"
        bold = (row['AUC_ensemble'] == best_auc and row['complete'])
        set_cell_font(cells[6], text, bold=bold, size=8)

    # ── Test de DeLong ──
    delong_img = REPORT_DIR / "delong_test.png"
    if delong_img.exists():
        doc.add_heading("Test de DeLong", level=2)
        doc.add_paragraph(
            "Test de DeLong par paires sur les probabilit\u00e9s moyenn\u00e9es. Les m\u00e9thodes avec un seul seed "
            "sont marqu\u00e9es N/A car la comparaison n'est pas fiable sans ex\u00e9cutions r\u00e9p\u00e9t\u00e9es."
        )
        doc.add_picture(str(delong_img), width=Inches(5.5))

    # ── Matrices de confusion ──
    cm_img = REPORT_DIR / "confusion_matrices.png"
    if cm_img.exists():
        doc.add_heading("Matrices de confusion", level=2)
        doc.add_picture(str(cm_img), width=Inches(6.0))

    # ── GradCAM ──
    gc_img = REPORT_DIR / "gradcam_examples.png"
    if gc_img.exists():
        doc.add_heading("GradCAM", level=2)
        doc.add_paragraph(
            "Visualisations GradCAM du seed avec le meilleur AUC (MLP Early Fusion). "
            "Les r\u00e9gions rouges indiquent une forte activation li\u00e9e \u00e0 la MA."
        )
        doc.add_picture(str(gc_img), width=Inches(6.0))

    # ── D\u00e9tails d'entra\u00eenement ──
    doc.add_heading("D\u00e9tails d'entra\u00eenement", level=2)
    details = [
        ("Backbone ResNet3D", "MONAI ResNet50, pr\u00e9-entra\u00een\u00e9 MedicalNet (23 jeux de donn\u00e9es)"),
        ("Fine-tuning", "30 \u00e9poques, gel\u00e9 3 premi\u00e8res \u00e9poques, LR diff\u00e9rentiel (backbone 10x plus faible)"),
        ("Pr\u00e9cision mixte", "AMP activ\u00e9, accumulation de gradient (2 \u00e9tapes, batch effectif=4)"),
        ("MLP tabulaire", "LayerNorm, couches cach\u00e9es [128, 64, 32], dropout=0.3, 100 \u00e9poques"),
        ("XGBoost tabulaire", "max_depth=6, lr=0.1, subsample=0.8, 300 rounds, early stopping=30"),
        ("D\u00e9s\u00e9quilibre de classes", "CrossEntropyLoss pond\u00e9r\u00e9e (78% CN / 22% MA)"),
        ("Early stopping", "Patience=20, min_epochs=30, moniteur : balanced accuracy"),
        ("Optimiseur", "AdamW avec cosine annealing + warmup (5 \u00e9poques)"),
        ("Seeds", f"{TARGET_SEEDS} seeds par m\u00e9thode (entra\u00eenement en cours pour certaines)"),
    ]
    t2 = doc.add_table(rows=len(details), cols=2)
    t2.style = "Light Grid Accent 1"
    for i, (comp, conf) in enumerate(details):
        set_cell_font(t2.rows[i].cells[0], comp, bold=True, size=8)
        set_cell_font(t2.rows[i].cells[1], conf, size=8)

    # Save
    out_path = REPORT_DIR / "resnet3d_fusion_report_multi_seed.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    generate_report()
